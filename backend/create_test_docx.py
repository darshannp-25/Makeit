from docx import Document

doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test document for conversion.')
doc.save('test_input.docx')
